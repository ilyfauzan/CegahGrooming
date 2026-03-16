import docx

doc_path = r'c:\Users\FAUZAN\Documents\skripsidraft.docx'
doc = docx.Document(doc_path)

# Content to insert
latar_belakang = [
    "1.1 Latar Belakang",
    "Di era digital, peningkatan pesat internet dan media sosial memiliki dua konsekuensi: akses informasi tak terbatas sekaligus bahaya besar bagi keselamatan anak, terutama eksploitasi dan kekerasan seksual online yang meningkat. Dengan ribuan kasus terdokumentasi, termasuk kekerasan seksual dan kejahatan siber yang melibatkan platform digital seperti WhatsApp, TikTok, Instagram, dan game online yang populer di kalangan remaja Indonesia, laporan tahunan KPAI mencatat tingkat kekerasan terhadap anak masih tinggi (KPAI, 2025).",
    "Data KPAI semakin menunjukkan bahwa grooming anak adalah modus operandi predator yang sulit dideteksi karena prosesnya bertahap: dimulai dengan pendekatan ramah, membangun kepercayaan, hingga mencapai konten seksual melalui percakapan panjang tanpa pengawasan orang tua atau algoritma platform (KPAI, 2026). Secara global, penelitian Buitrago et al. menemukan pola grooming yang khas, termasuk obrolan kecil tentang hobi atau sekolah, transisi ke pertanyaan personal (usia, alamat), dan permintaan foto intim. Pola-pola ini sering hilang jika dianalisis per pesan tunggal, menjadikan deteksi kontekstual penting untuk mencegah eksploitasi di Indonesia (Buitrago et al., 2022).",
    "Banyak penelitian telah mengembangkan sistem untuk mendeteksi predator seksual dan cyber grooming menggunakan machine learning dan Natural Language Processing (NLP). Early sexual predator detection (eSPD) didefinisikan pada dataset percakapan nyata oleh Vogt et al. Mereka menemukan bahwa model BERT dapat mendeteksi pola grooming sejak tahap awal dengan skor F1 yang tinggi. Namun, mereka tidak melakukan pengukuran risiko bertahap atau adaptasi multibahasa untuk klasifikasi biner (predator/non-predator) (Vogt et al., 2021).",
    "Untuk klasifikasi cybergrooming dalam konteks perlindungan anak online, Isaza et al. mengusulkan model hybrid CNN-NLP. Model ini dapat dengan akurat mengidentifikasi pesan berisiko tetapi menghasilkan false positive tinggi karena analisis level pesan tunggal mengabaikan konteks rangkaian dialog yang panjang (Isaza et al., 2022). Selain itu, penelitian yang dilakukan oleh Fatma et al. menunjukkan keunggulan transformer dalam menangkap nuansa bahasa halus, pola sarkasme, dan konteks budaya. Ini memiliki potensi besar untuk grooming yang rumit dan manipulatif, tetapi belum diintegrasikan secara optimal untuk situasi real-time di platform lokal Indonesia (Fatma et al., 2021).",
    "Oleh karena itu, skripsi ini menyarankan pembuatan dan penerapan Sistem Deteksi Dini Potensi Predator Seksual Online berbasis kecerdasan buatan yang menggunakan model BERT, metode analisis kontekstual percakapan, dan arsitektur hybrid scoring untuk mengintegrasikan embedding semantik dari setiap pesan: (1) kemungkinan klasifikasi grooming per pesan, (2) pola urutan dialog (membangun hubungan -> seksualisasi), dan (3) intensitas atau frekuensi ujaran berisiko. Skor risiko komprehensif ini dapat disesuaikan untuk peringatan dini real-time (Isaza et al., 2022; Vogt et al., 2021)."
]

rumusan_masalah = [
    "1.2 Rumusan Masalah",
    "Berdasarkan latar belakang tersebut, rumusan masalah penelitian ini adalah sebagai berikut:",
    "1. Bagaimana merancang dan mengimplementasikan model bahasa BERT (Bidirectional Encoder Representations from Transformers) dengan mekanisme sliding window untuk membangun sistem deteksi dini online grooming yang mampu memahami alur percakapan secara kontekstual?",
    "2. Sejauh mana efektivitas penggunaan metode Hybrid Scoring yang menggabungkan analisis kalimat tunggal (standalone) dan analisis konteks percakapan dalam meningkatkan akurasi identifikasi pola predator seksual dibandingkan dengan metode klasifikasi teks konvensional?",
    "3. Bagaimana akurasi kategorisasi tingkat risiko (Normal, Warning, Danger) yang dihasilkan oleh sistem dalam memberikan peringatan dini yang presisi terhadap tahapan-tahapan modulasi perilaku predator seksual online?"
]

tujuan_penelitian = [
    "1.3 Tujuan Penelitian",
    "Berdasarkan rumusan masalah di atas, tujuan dari penelitian ini adalah:",
    "1. Merancang dan mengimplementasikan model BERT dengan mekanisme sliding window untuk deteksi grooming yang peka terhadap konteks percakapan.",
    "2. Mengukur dan menganalisis efektivitas metode Hybrid Scoring dalam meningkatkan akurasi identifikasi predator seksual dibandingkan metode konvensional.",
    "3. Mengevaluasi tingkat akurasi klasifikasi risiko (Normal, Warning, Danger) pada setiap tahapan eskalasi obrolan predator."
]

landasan_teori = [
    "1.4 Landasan Teori",
    "",
    "1.4.1 Online Grooming",
    "Online grooming didefinisikan sebagai upaya sistematis yang dilakukan oleh pelaku dewasa untuk membangun hubungan emosional dan kepercayaan dengan anak atau remaja melalui platform digital (media sosial, aplikasi pesan, atau game online) dengan tujuan eksploitasi seksual. Proses ini bersifat progresif dan manipulatif, yang menurut Buitrago et al. (2022) sering kali melibatkan tahapan pengumpulan informasi personal, pembentukan rahasia bersama, hingga fase isolasi sosial terhadap korban.",
    "",
    "1.4.2 Natural Language Processing (NLP)",
    "Pemrosesan Bahasa Alami atau Natural Language Processing (NLP) merupakan cabang kecerdasan buatan yang berfokus pada interaksi antara komputer dan bahasa manusia. Dalam konteks deteksi grooming, NLP digunakan untuk melakukan ekstraksi fitur semantik dari teks percakapan, memungkinkan sistem untuk mengenali pola linguistik yang mencurigakan melalui teknik-teknik seperti tokenization, lemmatization, dan vector representation.",
    "",
    "1.4.3 Model BERT (Bidirectional Encoder Representations from Transformers)",
    "BERT merupakan arsitektur deep learning berbasis Transformer yang dilatih untuk memahami konteks sebuah kata berdasarkan posisi kata-kata sebelum dan sesudahnya secara simultan (bidirectional). Keunggulan utama BERT terletak pada mekanisme Self-Attention, yang memungkinkannya menangkap nuansa bahasa yang halus dan hubungan dependensi jangka panjang dalam sebuah kalimat dengan akurasi yang lebih tinggi dibandingkan model tradisional.",
    "",
    "1.4.4 Klasifikasi Teks Kontekstual",
    "Klasifikasi teks dalam penelitian ini melampaui metode tradisional yang hanya berbasis kata kunci tunggal. Pendekatan kontekstual melibatkan analisis terhadap intensi di balik pesan, di mana label klasifikasi ditentukan tidak hanya oleh konten literal pesan tersebut, tetapi juga oleh bagaimana pesan tersebut berkontribusi pada narasi percakapan secara keseluruhan.",
    "",
    "1.4.5 Mekanisme Sliding Window",
    "Mekanisme Sliding Window diimplementasikan untuk mengatasi keterbatasan memori pada model bahasa besar serta untuk menangkap pola kronologis dalam percakapan. Dengan menganalisis blok pesan terakhir (misalnya 10 pesan), sistem dapat mendeteksi \"eskalasi\" atau perubahan suhu percakapan dari tahap normal menuju tahap manipulatif secara dinamis.",
    "",
    "1.4.6 Machine Translation & Cross-Lingual Bridge",
    "Dikarenakan keterbatasan dataset grooming dalam Bahasa Indonesia, skema Translation-Bridge digunakan untuk menerjemahkan input pengguna ke dalam Bahasa Inggris sebelum diproses oleh model BERT yang telah di-fine-tune pada dataset internasional. Metode ini memungkinkan pemanfaatan transfer data lintas bahasa untuk meningkatkan performa deteksi pada bahasa-bahasa dengan sumber daya rendah (low-resource languages).",
    "",
    "1.4.7 Dataset PAN Sexual Predator Identification",
    "Dataset PAN merupakan standarisasi dataset internasional yang berisi ribuan log percakapan nyata yang melibatkan predator seksual dan korban. Dataset ini menjadi fondasi utama dalam melatih model untuk mengenali perilaku predator, karena mencakup variasi pola interaksi mulai dari obrolan kasual hingga fase eksploitasi yang eksplisit.",
    "",
    "1.4.8 Arsitektur Hybrid Scoring",
    "Arsitektur Hybrid Scoring menggabungkan dua nilai probabilitas deteksi: skor dari pesan tunggal (standalone score) dan skor dari alur percakapan (context score). Penggabungan ini bertujuan untuk meminimalisir angka False Positive (salah deteksi) dan memastikan bahwa peringatan bahaya diberikan hanya jika terdapat konsistensi pola predator baik secara tekstual maupun situasional."
]

full_new_chapter_1 = [
    "BAB I",
    "PENDAHULUAN"
] + latar_belakang + [""] + rumusan_masalah + [""] + tujuan_penelitian + [""] + landasan_teori

def find_paragraph_index(doc, marker, start_from=0):
    for i in range(start_from, len(doc.paragraphs)):
        if marker in doc.paragraphs[i].text:
            return i
    return -1

start_idx = find_paragraph_index(doc, "BAB  I  PENDAHULUAN")
if start_idx == -1:
    start_idx = find_paragraph_index(doc, "BAB I")

end_idx = find_paragraph_index(doc, "BAB II", start_from=start_idx+1 if start_idx != -1 else 0)

if start_idx != -1 and end_idx != -1:
    # Get the paragraph after the section
    after_p = doc.paragraphs[end_idx]
    
    # Delete from start to end-1
    for _ in range(end_idx - start_idx):
        p_to_del = doc.paragraphs[start_idx]
        p_to_del._element.getparent().remove(p_to_del._element)
    
    # Insert new lines
    for line in reversed(full_new_chapter_1):
        after_p.insert_paragraph_before(line)
        
    doc.save(doc_path)
    print("Surgical fix applied successfully.")
else:
    print(f"Failed to find markers. Start: {start_idx}, End: {end_idx}")

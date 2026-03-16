import docx
import os

doc_path = r'c:\Users\FAUZAN\Documents\skripsidraft.docx'
output_path = r'c:\Users\FAUZAN\Documents\skripsidraft.docx'

doc = docx.Document(doc_path)

def replace_section(doc, start_marker, end_marker, new_content):
    start_idx = -1
    end_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        if start_marker in p.text:
            start_idx = i
        if end_marker in p.text and start_idx != -1:
            end_idx = i
            break
            
    if start_idx != -1 and end_idx != -1:
        # Clear existing paragraphs between markers
        # We delete from end to start to avoid index shifting issues
        for _ in range(end_idx - start_idx - 1):
            p_to_del = doc.paragraphs[start_idx + 1]
            p_to_del._element.getparent().remove(p_to_del._element)
            
        # Add new content
        # Insert before the end_marker paragraph
        end_p = doc.paragraphs[start_idx + 1]
        for line in reversed(new_content.split('\n')):
            new_p = doc.paragraphs[start_idx].insert_paragraph_before(line)
            # Move it to the right place - actually insert_paragraph_before is not ideal for multiple lines
            # Let's just append after start_marker and then move them? No.
            
    else:
        print(f"Markers {start_marker} or {end_marker} not found.")

# Alternative safer way: 
def update_docx(doc_path):
    doc = docx.Document(doc_path)
    
    # Define sections
    rm_content = [
        "Rumusan Masalah",
        "Berdasarkan latar belakang tersebut, rumusan masalah penelitian ini adalah sebagai berikut:",
        "1. Bagaimana merancang dan mengimplementasikan model bahasa BERT (Bidirectional Encoder Representations from Transformers) dengan mekanisme sliding window untuk membangun sistem deteksi dini online grooming yang mampu memahami alur percakapan secara kontekstual?",
        "2. Sejauh mana efektivitas penggunaan metode Hybrid Scoring yang menggabungkan analisis kalimat tunggal (standalone) dan analisis konteks percakapan dalam meningkatkan akurasi identifikasi pola predator seksual dibandingkan dengan metode klasifikasi teks konvensional?",
        "3. Bagaimana akurasi kategorisasi tingkat risiko (Normal, Warning, Danger) yang dihasilkan oleh sistem dalam memberikan peringatan dini yang presisi terhadap tahapan-tahapan modulasi perilaku predator seksual online?"
    ]
    
    tujuan_content = [
        "Tujuan Penelitian",
        "Berdasarkan rumusan masalah di atas, tujuan dari penelitian ini adalah:",
        "1. Merancang dan mengimplementasikan model BERT dengan mekanisme sliding window untuk deteksi grooming yang peka terhadap konteks percakapan.",
        "2. Mengukur dan menganalisis efektivitas metode Hybrid Scoring dalam meningkatkan akurasi identifikasi predator seksual dibandingkan metode konvensional.",
        "3. Mengevaluasi tingkat akurasi klasifikasi risiko (Normal, Warning, Danger) pada setiap tahapan eskalasi obrolan predator."
    ]
    
    lt_content = [
        "1.4 Landasan Teori",
        "",
        "1.4.1 Online Grooming",
        "Online grooming didefinisikan sebagai upaya sistematis yang dilakukan oleh pelaku dewasa untuk membangun hubungan emosional dan kepercayaan dengan anak atau remaja melalui platform digital (media sosial, aplikasi pesan, atau game online) dengan tujuan eksploitasi seksual. Proses ini bersifat progresif dan manipulatif, yang menurut Buitrago et al. (2022) sering kali melibatkan tahapan pengumpulan informasi personal, pembentukan rahasia bersama, hingga fase isolasi sosial terhadap korban.",
        "",
        "1.4.2 Natural Language Processing (NLP)",
        "Pemrosesan Bahasa Alami atau Natural Language Processing (NLP) merupakan cabang kecerdasan buatan yang berfokus pada interaksi antara komputer dan bahasa manusia. Dalam konteks deteksi grooming, NLP digunakan untuk melakukan ekstraksi fitur semantik dari teks percakapan, memungkinkan sistem untuk mengenali pola linguistik yang mencurigakan melalui teknik-teknik seperti tokenization, lemmatization, dan vector representation.",
        "",
        "1.4.3 Model BERT (Bidirectional Encoder Representations from Transformers)",
        "BERT merupakan arsitektur deep learning berbasis Transformer yang dilatih untuk memahami konteks sebuah kata berdasarkan posisi kata-kata sebelum dan sesudahnya secara simultan (bidirectional). Keunggulan utama BERT terletak pada mekanisme Self-Attention, yang memungkinkannya menangkap nuansa bahasa yang halus dan hubungan dependensi jangka panjang dalam sebuah kalimat dengan akurasi yang lebih tinggi dibandingkan model tradisional.",
        "",
        "1.4.4 Klasifikasi Teks Kontekstual",
        "Klasifikasi teks dalam penelitian ini melampaui metode tradisional yang hanya berbasis kata kunci tunggal. Pendekatan kontekstual melibatkan analisis terhadap intensi di balik pesan, di mana label klasifikasi ditentukan tidak hanya oleh konten literal pesan tersebut, tetapi juga oleh bagaimana pesan tersebut berkontribusi pada narasi percakapan secara keseluruhan.",
        "",
        "1.4.5 Mekanisme Sliding Window",
        "Mekanisme Sliding Window diimplementasikan untuk mengatasi keterbatasan memori pada model bahasa besar serta untuk menangkap pola kronologis dalam percakapan. Dengan menganalisis blok pesan terakhir (misalnya 10 pesan), sistem dapat mendeteksi \"eskalasi\" atau perubahan suhu percakapan dari tahap normal menuju tahap manipulatif secara dinamis.",
        "",
        "1.4.6 Machine Translation & Cross-Lingual Bridge",
        "Dikarenakan keterbatasan dataset grooming dalam Bahasa Indonesia, skema Translation-Bridge digunakan untuk menerjemahkan input pengguna ke dalam Bahasa Inggris sebelum diproses oleh model BERT yang telah di-fine-tune pada dataset internasional. Metode ini memungkinkan pemanfaatan transfer data lintas bahasa untuk meningkatkan performa deteksi pada bahasa-bahasa dengan sumber daya rendah (low-resource languages).",
        "",
        "1.4.7 Dataset PAN Sexual Predator Identification",
        "Dataset PAN merupakan standarisasi dataset internasional yang berisi ribuan log percakapan nyata yang melibatkan predator seksual dan korban. Dataset ini menjadi fondasi utama dalam melatih model untuk mengenali perilaku predator, karena mencakup variasi pola interaksi mulai dari obrolan kasual hingga fase eksploitasi yang eksplisit.",
        "",
        "1.4.8 Arsitektur Hybrid Scoring",
        "Arsitektur Hybrid Scoring menggabungkan dua nilai probabilitas deteksi: skor dari pesan tunggal (standalone score) dan skor dari alur percakapan (context score). Penggabungan ini bertujuan untuk meminimalisir angka False Positive (salah deteksi) dan memastikan bahwa peringatan bahaya diberikan hanya jika terdapat konsistensi pola predator baik secara tekstual maupun situasional."
    ]

    def find_and_replace_range(doc, start_key, end_key, content_list):
        paragraphs = list(doc.paragraphs)
        start_idx = -1
        end_idx = -1
        
        for i, p in enumerate(paragraphs):
            if start_key in p.text:
                start_idx = i
            if end_key in p.text and start_idx != -1:
                end_idx = i
                break
        
        if start_idx != -1 and end_idx != -1:
            # Delete content in between
            for _ in range(end_idx - start_idx - 1):
                p_to_remove = doc.paragraphs[start_idx + 1]
                p_to_remove._element.getparent().remove(p_to_remove._element)
            
            # Insert new paragraphs
            # We must be careful about where we insert.
            # Insert after start_idx
            current_p = doc.paragraphs[start_idx]
            for line in content_list[1:]: # Skip the header if it stayed
                new_p = current_p.insert_paragraph_before(line) # This actually inserts BEFORE
                # We need a different logic to insert AFTER.
                # Actually, iterate through the list and insert before end_idx?
            
            # Let's use a simpler logic: replace the text of paragraphs if available, else add.
            # RE-write:
            
            # 1. Store the paragraph after section
            after_p = doc.paragraphs[start_idx + 1] 
            # 2. Delete everything from start_idx to end_idx-1
            for _ in range(end_idx - start_idx):
                p_to_remove = doc.paragraphs[start_idx]
                p_to_remove._element.getparent().remove(p_to_remove._element)
            
            # 3. Insert new paragraphs BEFORE after_p
            for line in reversed(content_list):
                after_p.insert_paragraph_before(line)
            return True
        return False

    # Perform updates
    find_and_replace_range(doc, "Rumusan Masalah", "Tujuan Penelitian", rm_content)
    # Reload paras after change
    find_and_replace_range(doc, "Tujuan Penelitian", "Manfaat Penelitian", tujuan_content)
    find_and_replace_range(doc, "1.4 Landasan Teori", "BAB II  METODOLOGI PENELITIAN", lt_content)
    
    doc.save(doc_path)
    print("Document updated successfully.")

update_docx(doc_path)

import docx

doc_path = r'c:\Users\FAUZAN\Documents\skripsidraft.docx'
doc = docx.Document(doc_path)

# Restore content to placeholders
latar_belakang = [
    "1.1 Latar Belakang",
    "Di era digital, peningkatan pesat internet dan media sosial memiliki dua konsekuensi: akses informasi tak terbatas sekaligus bahaya besar bagi keselamatan anak, terutama eksploitasi dan kekerasan seksual online yang meningkat. Dengan ribuan kasus terdokumentasi, termasuk kekerasan seksual dan kejahatan siber yang melibatkan platform digital seperti WhatsApp, TikTok, Instagram, dan game online yang populer di kalangan remaja Indonesia, laporan tahunan KPAI mencatat tingkat kekerasan terhadap anak masih tinggi (KPAI, 2025).",
    "Data KPAI semakin menunjukkan bahwa grooming anak adalah modus operandi predator yang sulit dideteksi karena prosesnya bertahap: dimulai dengan pendekatan ramah, membangun kepercayaan, hingga mencapai konten seksual melalui percakapan panjang tanpa pengawasan orang tua atau algoritma platform (KPAI, 2026). Secara global, penelitian Buitrago et al. menemukan pola grooming yang khas, termasuk obrolan kecil tentang hobi atau sekolah, transisi ke pertanyaan personal (usia, alamat), dan permintaan foto intim. Pola-pola ini sering hilang jika dianalisis per pesan tunggal, menjadikan deteksi kontekstual penting untuk mencegah eksploitasi di Indonesia (Buitrago et al., 2022).",
    "Banyak penelitian telah mengembangkan sistem untuk mendeteksi predator seksual dan cyber grooming menggunakan machine learning dan Natural Language Processing (NLP). Early sexual predator detection (eSPD) didefinisikan pada dataset percakapan nyata oleh Vogt et al. Mereka menemukan bahwa model BERT dapat mendeteksi pola grooming sejak tahap awal dengan skor F1 yang tinggi. Namun, mereka tidak melakukan pengukuran risiko bertahap atau adaptasi multibahasa untuk klasifikasi biner (predator/non-predator) (Vogt et al., 2021).",
    "Untuk klasifikasi cybergrooming dalam konteks perlindungan anak online, Isaza et al. mengusulkan model hybrid CNN-NLP. Model ini dapat dengan akurat mengidentifikasi pesan berisiko tetapi menghasilkan false positive tinggi karena analisis level pesan tunggal mengabaikan konteks rangkaian dialog yang panjang (Isaza et al., 2022). Selain itu, penelitian yang dilakukan oleh Fatma et al. menunjukkan keunggulan transformer dalam menangkap nuansa bahasa halus, pola sarkasme, dan konteks budaya. Ini memiliki potensi besar untuk grooming yang rumit and manipulatif, tetapi belum diintegrasikan secara optimal untuk situasi real-time di platform lokal Indonesia (Fatma et al., 2021).",
    "Oleh karena itu, skripsi ini menyarankan pembuatan dan penerapan Sistem Deteksi Dini Potensi Predator Seksual Online berbasis kecerdasan buatan yang menggunakan model BERT, metode analisis kontekstual percakapan, and arsitektur hybrid scoring untuk mengintegrasikan embedding semantik dari setiap pesan: (1) kemungkinan klasifikasi grooming per pesan, (2) pola urutan dialog (membangun hubungan -> seksualisasi), and (3) intensitas atau frekuensi ujaran berisiko. Skor risiko komprehensif ini dapat disesuaikan untuk peringatan dini real-time (Isaza et al., 2022; Vogt et al., 2021)."
]

rumusan_masalah = [
    "1.2 Rumusan Masalah",
    "Di era digital saat ini, peningkatan akses internet dan media sosial di Indonesia telah memicu lonjakan kasus eksploitasi seksual online terhadap anak, dengan grooming sebagai modus utama yang sulit dideteksi karena sifatnya bertahap dan kontekstual.",
    "Berdasarkan latar belakang tersebut, rumusan masalah penelitian ini adalah sebagai berikut:",
    "Bagaimana merancang dan mengimplementasikan model bahasa BERT (Bidirectional Encoder Representations from Transformers) dengan mekanisme sliding window untuk membangun sistem deteksi dini online grooming yang mampu memahami alur percakapan secara kontekstual?",
    "Sejauh mana efektivitas penggunaan metode Hybrid Scoring yang menggabungkan analisis kalimat tunggal (standalone) dan analisis konteks percakapan dalam meningkatkan akurasi identifikasi pola predator seksual dibandingkan dengan metode klasifikasi teks konvensional?"
]

tujuan_penelitian = [
    "1.3 Tujuan Penelitian",
    "Merancang dan mengimplementasikan model BERT with sliding window untuk deteksi grooming kontekstual.",
    "Mengukur efektivitas Hybrid Scoring dalam meningkatkan akurasi dibanding metode konvensional."
]

landasan_teori = [
    "1.4 Landasan Teori",
    "1.4.1 Subjudul 1",
    "1.4.2 Subjudul 2",
    "1.4.3 Subjudul 3"
]

full_revert_content = [
    "BAB I",
    "PENDAHULUAN"
] + latar_belakang + [""] + rumusan_masalah + [""] + tujuan_penelitian + [""] + landasan_teori

def find_paragraph_index(doc, marker, start_from=0):
    for i in range(start_from, len(doc.paragraphs)):
        if marker in doc.paragraphs[i].text:
            return i
    return -1

start_idx = find_paragraph_index(doc, "BAB I")
if start_idx == -1:
    start_idx = find_paragraph_index(doc, "BAB  I")

end_idx = find_paragraph_index(doc, "BAB II", start_from=start_idx+1 if start_idx != -1 else 0)

if start_idx != -1 and end_idx != -1:
    after_p = doc.paragraphs[end_idx]
    
    # Delete current Chapter 1
    for _ in range(end_idx - start_idx):
        p_to_del = doc.paragraphs[start_idx]
        p_to_del._element.getparent().remove(p_to_del._element)
    
    # Insert revert content IN ORDER
    for line in full_revert_content[::-1]:
        after_p.insert_paragraph_before(line)
        
    doc.save(doc_path)
    print("Reverted to normal state successfully.")
else:
    print(f"Failed to find markers. Start: {start_idx}, End: {end_idx}")
